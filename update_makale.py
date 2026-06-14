import docx

docx_path = r'C:\Users\Lenovo\Desktop\1\YAZILIM MÜHENDİSLİĞİ\MAKALE.docx'

try:
    doc = docx.Document(docx_path)
    
    insert_after_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "2. ORM Entegrasyonu:" in p.text:
            insert_after_idx = i
            break
            
    if insert_after_idx != -1:
        # Create a new paragraph after index
        p_ref = doc.paragraphs[insert_after_idx]
        new_text = "3. İlişkisel Bütünlük (Referential Integrity) ve Kaskad: Veritabanındaki Foreign Key bağlarının tutarlı kalması için ORM seviyesinde CascadeType.ALL tanımlanmış, böylece bir etkinlik silinirken ona ait başvuruların (kayıtların) otomatik olarak silinmesi sağlanarak bütünlük ihlali ve hata durumları engellenmiştir."
        # Using the same style
        new_p = p_ref.insert_paragraph_before("")
        p_ref._element.addprevious(new_p._element) # wait, insert_paragraph_before puts it before, but we want it after.
        # Actually it's easier to insert_before the NEXT paragraph.
        p_next = doc.paragraphs[insert_after_idx + 1]
        p_inserted = p_next.insert_paragraph_before(new_text)
        
        # Copy formatting (left indent, etc.)
        p_inserted.paragraph_format.left_indent = p_ref.paragraph_format.left_indent
        p_inserted.paragraph_format.first_line_indent = p_ref.paragraph_format.first_line_indent
        p_inserted.alignment = p_ref.alignment
        
        print("Successfully updated the Word document.")
        
    doc.save(docx_path)
    print("Saved MAKALE.docx")

except Exception as e:
    print(f"Error: {e}")
