import os
from docx import Document

docx_path = r'C:\Users\Lenovo\Desktop\1\YAZILIM MÜHENDİSLİĞİ\Conference-template-A4-Guncel.docx'

try:
    doc = Document(docx_path)

    # Append new section for the UI and localization changes
    doc.add_heading('Nihai Sistem İyileştirmeleri ve Revizyonlar', level=2)

    p1 = doc.add_paragraph()
    p1.add_run('Son iterasyonda Actimatter platformuna şu eklemeler ve iyileştirmeler yapılmıştır:').bold = True

    doc.add_paragraph(
        '- Çoklu Dil Desteği (i18n): Sisteme Türkçe (TR) ve İngilizce (EN) olmak üzere çift dil desteği eklenmiştir. '
        'Kullanıcılar navbar üzerinden diller arası geçiş yapabilmektedir. Tüm arayüz metinleri ("Pending", "Approved" gibi terimler dahil) '
        'ilgili dil dosyalarından okunarak çoklu dil prensibine uygun hale getirilmiştir.'
    )

    doc.add_paragraph(
        '- Dinamik Renk Paleti ve Tema Değişikliği: Varsayılan mavi ve çivit (indigo) renk paletinden vazgeçilerek, '
        'kullanıcının isteği doğrultusunda tüm arayüz bileşenleri "Yeşil" ve "Zümrüt Yeşili" (Emerald) renk tonlarına dönüştürülmüştür. '
        'Buton hover efektleri, navbar, badge ve kart arkaplanları yeşil tema uyumluluğuna göre güncellenmiştir.'
    )

    doc.add_paragraph(
        '- Etkinlik Detay Modal Ekranı: Katılımcıların etkinlik kartlarına "Devamını gör..." şeklinde basarak, o etkinliğin detaylı '
        'açıklamasına, konumuna, kapasite durumuna ulaşıp aynı ekran üzerinden kayıt olabilecekleri dinamik bir "Event Details Modal" yapısı inşa edilmiştir.'
    )

    doc.add_paragraph(
        '- Gelişmiş Tekrar Kayıt Kontrolü (Idempotency): Bir kullanıcının aynı etkinliğe sadece bir kere başvurabilmesi sağlanmıştır. '
        'Daha önce başvurulan etkinliklerin "Hemen Kayıt Ol" butonları "Kayıt Alındı" (Registration Received) olarak disable edilmekte '
        've sunucu tarafında existsByEventIdAndUserId kontrolüyle çift kayıt engellenmektedir.'
    )

    doc.save(docx_path)
    print(f"Successfully updated {docx_path}")
except Exception as e:
    print(f"Error updating docx: {e}")
