import docx

docx_path = r'C:\Users\Lenovo\Desktop\1\YAZILIM MÜHENDİSLİĞİ\Conference-template-A4-Guncel.docx'

try:
    doc = docx.Document(docx_path)
    
    # 1. Remove the chatty section at the end
    paragraphs = doc.paragraphs
    start_delete_idx = -1
    for i, p in enumerate(paragraphs):
        if "Nihai Sistem İyileştirmeleri ve Revizyonlar" in p.text:
            start_delete_idx = i
            break
            
    if start_delete_idx != -1:
        # Delete from start_delete_idx to end
        for i in range(len(paragraphs) - 1, start_delete_idx - 1, -1):
            p = paragraphs[i]
            p._element.getparent().remove(p._element)

    # 2. Add Mermaid Code under "Şekil 1."
    # Find "Şekil 1."
    mermaid_code = """
```mermaid
erDiagram
    USER {
        Long id PK
        String username
        String email
        String password
        String role
    }
    EVENT {
        Long id PK
        String title
        String description
        LocalDateTime eventDate
        Integer capacity
        String location
        Long coordinator_id FK
    }
    REGISTRATION {
        Long id PK
        String status
        LocalDateTime registrationDate
        Long event_id FK
        Long user_id FK
    }

    USER ||--o{ EVENT : "creates"
    USER ||--o{ REGISTRATION : "makes"
    EVENT ||--o{ REGISTRATION : "receives"
```
"""
    
    for i, p in enumerate(doc.paragraphs):
        if "Şekil 1." in p.text:
            new_p = p.insert_paragraph_before(mermaid_code.strip())
            break

    # 3. Insert Academic Features before "VIII. SONUÇ"
    academic_text = [
        "Sistemin işlevselliğini ve kullanıcı deneyimini artırmak amacıyla mimariye çeşitli iyileştirmeler entegre edilmiştir:",
        "1. Çoklu Dil Desteği (i18n): Uluslararasılaşma standartları gereği sisteme Türkçe ve İngilizce dil destekleri entegre edilmiş, arayüz bileşenleri dinamik dil değişimine uygun hale getirilmiştir.",
        "2. Tema ve Arayüz Optimizasyonu: Sistem arayüzünde kullanılan renk paleti, kullanıcı deneyimini artıracak şekilde 'Emerald' (Zümrüt Yeşili) tonlarına geçirilmiş ve görsel hiyerarşi yeniden düzenlenmiştir.",
        "3. Etkinlik Detay Modülü: Katılımcıların etkinlik bilgilerine daha hızlı erişebilmesi için, sayfa yenilenmeden asenkron olarak çalışan ve kayıt olma imkanı sunan bir modal pencere (dialog) yapısı geliştirilmiştir.",
        "4. Kayıt Tutarlılığı (Idempotency): Veri bütünlüğünü sağlamak adına, bir kullanıcının aynı etkinliğe birden fazla kez kaydolmasını engelleyen kısıtlamalar sunucu (backend) katmanında uygulanmıştır."
    ]

    for i, p in enumerate(doc.paragraphs):
        if "VIII. SONUÇ" in p.text:
            # Insert academic text before this paragraph
            for text in academic_text:
                inserted = p.insert_paragraph_before(text)
            break

    new_docx_path = r'C:\Users\Lenovo\Desktop\1\YAZILIM MÜHENDİSLİĞİ\Conference-template-A4-Guncel-v2.docx'
    doc.save(new_docx_path)
    print(f"Successfully cleaned and updated {new_docx_path}")

except Exception as e:
    print(f"Error: {e}")
